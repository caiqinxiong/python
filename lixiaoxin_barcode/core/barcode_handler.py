# -*- coding: utf-8 -*-
__author__ = 'caiqinxiong_cai'
#2019/12/17 10:44
import os
import docx
import barcode
from barcode.writer import ImageWriter
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT,WD_TAB_LEADER  #设置制表符等
from docx.shared import Inches   #设置图像大小
from docx.shared import Pt    #设置像素、缩进等
from docx.shared import RGBColor    #设置字体颜色
from docx.shared import Length    #设置宽度
from conf import settings as ss

def read_file():
    '''读取文本信息'''
    data_list = []
    with open(ss.INPUT_FILE,mode='r') as f:
        for line in f:
            line = line.strip()
            if line:data_list.append(line)
    return data_list

def create_barcode(content,bracode_name):
    '''条形码生成'''
    EAN = barcode.get_barcode_class(ss.CODE_FORMAT)  # 创建Code128格式的条形码格式对象
    ean = EAN(content, writer=ImageWriter())  # 创建条形码对象，内容为content参数传入
    bracode_name = os.path.join(ss.IMG_PAHT,bracode_name)
    return ean.save(bracode_name, {'write_text': False})  # 保存条形码图片，并返回保存路径。图片格式为png,{'write_text': False}为不要底下的文字内容

def get_barcode():
    '''获取条形码图片地址'''
    barcode_list = []
    for i in ss.BAR_NAME_LIST:
        name = i + '.png'
        name = os.path.join(ss.IMG_PAHT,name)
        barcode_list.append(name)
    return barcode_list


def write_docx(data_list):
    '''将条形码写入docx'''
    barcode_list = get_barcode() # 获取条形码图片地址
    document=docx.Document()   # 创建一个空白文档
    table=document.add_table(rows=4*6,cols=2*3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.styles['Normal'].font.name = '宋体'  # 设置西文字体
    document.styles['Normal'].font.size = Pt(7)  # 设置字号
    document.styles['Normal'].paragraph_format.space_after = Pt(0)  # 设置段后间距
    # document.styles['Normal'].paragraph_format.left_indent = Inches(0.1) # 设置缩进默认Inches(0.5)等于四个空格
    # 写入logo
    n = 0
    for i in range(6):
        m = 0
        for j in range(3):
            run = document.tables[0].cell(n, m).paragraphs[0].add_run()
            run.add_picture(ss.LOGO_PNG,width=Inches(0.8), height=Inches(0.2))

            for k in range(len(data_list)):
                cell = table.cell(k+n, 1+m)
                cell.text = ' ' + data_list[k] + '\n'
                run = document.tables[0].cell(k+n, 1+m).paragraphs[0].add_run()
                run.add_picture(barcode_list[k], width=Inches(1.0), height=Inches(0.2))
                document.tables[0].cell(k+n, 1+m).paragraphs[0].add_run()
            DATE = ' DATE:' + ss.DAY_TIME + '\n'
            cell = table.cell(3+n, 1+m)
            cell.text = DATE

            m = m+2
        n = n + 4

    document.save(ss.OUPUT_FILE)



def main():
    '''主逻辑'''
    # 1.读取文本内容
    data_list = read_file()
    
    # 2.循环生成条形码图片
    for i in range(len(data_list)):create_barcode(data_list[i],ss.BAR_NAME_LIST[i])

    # 3.将条形码插入docx
    write_docx(data_list)

main()