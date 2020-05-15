# -*- coding: utf-8 -*-
__author__ = 'caiqinxiong_cai'
#2019/12/17 10:44
# 解决pip下载速度慢的问题
# https://blog.csdn.net/wolflikeinnocence/article/details/80140064

import os
import docx
import barcode
import xlrd
from barcode.writer import ImageWriter
from docx.shared import Inches   #设置图像大小
from docx.shared import Pt    #设置像素、缩进等
from conf import settings as ss
from core.log import Log as log

def read_excel(start_row=ss.START_ROW, end_row=0, sheet_index=0):
    '''读取表格信息'''
    date_list = []
    workbook = xlrd.open_workbook(ss.INPUT_FILE)
    # 选择第1个sheet表，索引从0开始
    sheet = workbook.sheet_by_index(sheet_index)
    if not end_row:end_row = sheet.nrows
    # 从表格中读取数据，追加到列表中
    log.readAndWrite('从表格中读取数据，追加到列表中,跳过表头从%s行开始读取数据，共%s行数据！' % (start_row, sheet.nrows - start_row))
    for i in range(start_row,end_row):
        date_list.append(sheet.row_values(i))
    else:
        log.readAndWrite('数据读取完成！')
    return date_list

def create_barcode(content,bracode_name):
    '''条形码生成'''
    content = str(content)
    EAN = barcode.get_barcode_class(ss.CODE_FORMAT)  # 创建Code128格式的条形码格式对象
    ean = EAN(content, writer=ImageWriter())  # 创建条形码对象，内容为content参数传入
    bracode_name = os.path.join(ss.IMG_PAHT,bracode_name)
    log.readAndWrite('内容为%s的条形码%s生成ok' % (content,bracode_name))
    # 保存条形码图片，并返回保存路径。图片格式为png,{'write_text': False}为不要底下的文字内容,'quiet_zone'为调整图片左右白边大小，看源码
    # default_writer_options = { # 默认参数
    #     'module_width': 0.2,
    #     'module_height': 15.0,
    #     'quiet_zone': 6.5,
    #     'font_size': 10,
    #     'text_distance': 5.0,
    #     'background': 'white',
    #     'foreground': 'black',
    #     'write_text': True,
    #     'text': '',
    # }
    return ean.save(bracode_name, {'write_text': False, 'quiet_zone': 0})


def wirte_barcode(document,data_list,i,n,m):
    '''写入条形码'''
    # 写入logo
    table = document.tables[0].cell(n, m).paragraphs[0].add_run()
    table.add_picture(ss.LOGO_PNG, width=Inches(ss.LOGO_W), height=Inches(ss.LOGO_H))
    # 写入条形码
    table = document.tables[0].cell(n, m+1).paragraphs[0].add_run()
    table.text = ' PART NO.:' + str(data_list[i][0]) + '\n'
    name = str(i) + '_0.png'
    # PART_NO_W = ss.PART_NO_W * len(data_list[i][0]) * 100 / 13 * 0.01  # 根据字符长度决定条形码的长度，字符个数以13个为基线
    # if PART_NO_W < 0.3:PART_NO_W = 0.3 # 限制条形码最小长度
    # if PART_NO_W > 1.7:PART_NO_W = 1.7 # 限制条形码最大长度
    table.add_picture(os.path.join(ss.IMG_PAHT, name), width=Inches(ss.PART_NO_W), height=Inches(ss.PART_NO_H))
    # table.add_picture(os.path.join(ss.IMG_PAHT, name),height=Inches(ss.PART_NO_H))
    table = document.tables[0].cell(n, m+1).paragraphs[0].add_run()
    table.text = '\n PACK NO.:' + str(data_list[i][1]) + ' \n'
    name = str(i) + '_1.png'
    # PACK_NO_W = ss.PACK_NO_W * len(data_list[i][1])*100/13*0.01 # 根据字符长度决定条形码的长度，字符个数以13个为基线
    # if PACK_NO_W < 0.3:PACK_NO_W = 0.3 # 限制条形码最小长度
    # if PACK_NO_W > 1.7:PACK_NO_W = 1.7 # 限制条形码最大长度
    table.add_picture(os.path.join(ss.IMG_PAHT, name), width=Inches(ss.PACK_NO_W), height=Inches(ss.PACK_NO_H))
    # table.add_picture(os.path.join(ss.IMG_PAHT, name), height=Inches(ss.PACK_NO_H))
    table = document.tables[0].cell(n, m+1).paragraphs[0].add_run()
    table.text = '\n Quantity:' + str(data_list[i][2]).split('.')[0] + '\n'
    name = str(i) + '_2.png'
    # QUANTITY_W = ss.QUANTITY_W * data_list[i][2] *100/50*0.01 # 根据数值大小决定条形码长度，以50为基线
    # if QUANTITY_W < 0.3:QUANTITY_W = 0.3 # 限制条形码最小长度
    # if QUANTITY_W > 1.7:QUANTITY_W = 1.7 # 限制条形码最大长度
    table.add_picture(os.path.join(ss.IMG_PAHT, name), width=Inches(ss.QUANTITY_W), height=Inches(ss.QUANTITY_H))
    # table.add_picture(os.path.join(ss.IMG_PAHT, name),  height=Inches(ss.QUANTITY_H))
    DATE = '\n ' + 'DATE:' + ss.DAY_TIME
    table = document.tables[0].cell(n, m+1).paragraphs[0].add_run()
    table.text = DATE  # 写入当前日期

def write_docx(data_list):
    '''将条形码写入docx'''
    document=docx.Document(ss.TEM_DOC)   # 打开模板
    document.styles['Normal'].font.name = ss.WORD_FORM  # 设置西文字体
    document.styles['Normal'].font.size = Pt(ss.WORD_SIZE)  # 设置字号
    document.styles['Normal'].paragraph_format.space_after = Pt(0)  # 设置段后间距
    n = 0;m = 0
    for i in range(len(data_list)):
        # 写入条形码
        wirte_barcode(document, data_list, i, n, m)
        m = m+2
        if m > ss.BARCODE_NUM+1:m = 0;n = n+1
    else:
        log.readAndWrite('文档内容填写完成，地址为：\n%s' % ss.OUPUT_FILE)

    # 保存docx文档
    return  document.save(ss.OUPUT_FILE)
    

